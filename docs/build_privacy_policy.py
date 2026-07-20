from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "docs/X-Post-History-Privacy-Policy.docx"
BLUE = "1D9BF0"
DARK = "0F1419"
MUTED = "536471"
LIGHT = "E8F5FD"
BORDER = "CFD9DE"


def set_font(run, name="Calibri", size=11, bold=False, color=DARK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("X POST HISTORY  |  PRIVACY"), size=8.5, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("めっちゃX  •  プライバシーポリシー"), size=8.5, color=MUTED)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(4)
set_font(kicker.add_run("CHROME EXTENSION  •  USER DATA POLICY"), size=9, bold=True, color=BLUE)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(5)
set_font(title.add_run("プライバシーポリシー"), size=25, bold=True, color=DARK)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_font(subtitle.add_run("めっちゃX（X Post History）"), size=13, bold=True, color=MUTED)

meta = doc.add_table(rows=2, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.style = "Table Grid"
set_table_geometry(meta, [2700, 6660])
mark_header_row(meta.rows[0])
for row, (label, value) in zip(meta.rows, (("施行日", "2026年7月20日"), ("対象", "Chrome拡張機能「めっちゃX」"))):
    shade(row.cells[0], LIGHT)
    set_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color="1F4D78")
    set_font(row.cells[1].paragraphs[0].add_run(value))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.line_spacing = 1.25
set_font(p.add_run("要点  "), bold=True, color=BLUE)
set_font(p.add_run("本拡張機能が扱う履歴データは、お使いのChromeブラウザ内にのみ保存されます。外部サーバーへの送信、第三者への提供、販売、広告利用または分析利用は行いません。"), bold=True)

doc.add_heading("1. 適用範囲", level=1)
add_body(doc, "本プライバシーポリシーは、Chrome拡張機能「めっちゃX」（以下「本拡張機能」）におけるユーザーデータの取扱いに適用されます。本拡張機能は、X（x.com）上で表示されたポストを利用者自身が後から検索・閲覧できるようにすることを単一の目的としています。")

doc.add_heading("2. 取り扱うデータ", level=1)
add_body(doc, "本拡張機能は、利用者が初回同意画面で明示的に同意した場合に限り、次の情報を処理し、Chromeのローカルストレージに保存します。")
for item in (
    "ポストの識別子、URL、投稿日時",
    "ポスト本文、投稿者の表示名およびユーザー名",
    "ポストに含まれる画像URLおよび動画サムネイルURL",
    "ポストを最初および最後に表示した日時、表示回数、画面内での累計表示時間",
    "利用者がお気に入りに設定した状態",
    "履歴記録の有効・停止状態および同意状態などの拡張機能設定",
):
    add_bullet(doc, item)
add_body(doc, "本拡張機能は、Xのパスワード、認証Cookie、ダイレクトメッセージ、決済情報、健康情報、正確な位置情報を取得する目的では設計されておらず、これらを保存しません。")

doc.add_heading("3. データの利用目的", level=1)
add_body(doc, "取り扱うデータは、次の利用者向け機能を提供するためにのみ使用します。")
for item in (
    "X上で表示されたポストの閲覧履歴を一覧表示すること",
    "キーワード、日付、表示時間、お気に入り状態による履歴検索および絞り込み",
    "閲覧回数、累計表示時間、画像・動画サムネイルの表示",
    "利用者による履歴の個別削除、一括削除およびJSON形式での書き出し",
):
    add_bullet(doc, item)

doc.add_heading("4. 保存場所・外部送信・第三者提供", level=1)
add_body(doc, "保存場所：", bold_lead="保存場所：")
add_body(doc, "履歴と設定は、Chrome拡張機能の chrome.storage.local を使用して利用者のブラウザ内に保存されます。保存件数の上限は3,000件です。")
add_body(doc, "外部送信：", bold_lead="外部送信：")
add_body(doc, "本拡張機能は、保存した履歴・設定・閲覧行動を、開発者のサーバーまたは第三者のサーバーへ送信しません。解析サービス、広告サービス、トラッキングサービスも使用しません。")
add_body(doc, "第三者提供：", bold_lead="第三者提供：")
add_body(doc, "ユーザーデータを販売、貸与、共有または第三者へ提供しません。信用力の判断、融資、広告、リターゲティングその他の目的にも使用しません。")

doc.add_heading("5. 権限の利用", level=1)
permissions = doc.add_table(rows=1, cols=2)
permissions.style = "Table Grid"
permissions.alignment = WD_TABLE_ALIGNMENT.LEFT
permissions.rows[0].cells[0].text = "権限"
permissions.rows[0].cells[1].text = "利用理由"
for cell in permissions.rows[0].cells:
    shade(cell, "E8EEF5")
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, color="1F4D78")
mark_header_row(permissions.rows[0])
for left, right in (
    ("storage", "閲覧履歴、お気に入り、表示時間、記録設定および同意状態を利用者のブラウザ内に保存するため。"),
    ("https://x.com/*", "X上に表示されたポストの情報を履歴として記録し、Xのサイドバーに履歴画面を追加するため。アクセス対象はx.comのみに限定しています。"),
):
    cells = permissions.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(left), bold=True)
    set_font(cells[1].paragraphs[0].add_run(right))
set_table_geometry(permissions, [2700, 6660])

doc.add_heading("6. リモートコード", level=1)
add_body(doc, "本拡張機能はリモートコードを使用しません。実行されるJavaScriptおよびCSSはすべて拡張機能パッケージ内に含まれています。外部スクリプト、外部Wasm、eval() または文字列から生成したコードの実行は行いません。")

doc.add_heading("7. 利用者による管理と削除", level=1)
for item in (
    "初回同意前は履歴を記録しません。",
    "履歴画面から記録を一時停止または再開できます。再開時には同意が必要です。",
    "履歴は個別または一括で削除できます。",
    "拡張機能をアンインストールすると、Chromeが管理する本拡張機能のローカルデータは削除されます。",
    "必要に応じて履歴をJSON形式で利用者自身の端末へ書き出せます。",
):
    add_bullet(doc, item)

doc.add_heading("8. データの保持期間", level=1)
add_body(doc, "履歴データは、利用者が削除するか、本拡張機能をアンインストールするまでブラウザ内に保持されます。保存上限を超えた場合は、古い履歴から順に保存対象外となります。外部サーバーにはデータを保持しません。")

doc.add_heading("9. セキュリティ", level=1)
add_body(doc, "本拡張機能は、要求する権限とアクセス先を単一目的に必要な最小範囲に限定します。ユーザーデータをネットワーク送信しない設計とし、Chromeが提供する拡張機能専用ストレージを使用します。利用者は端末およびChromeプロフィールへのアクセスを適切に管理してください。")

doc.add_heading("10. Chrome ウェブストアの Limited Use 要件", level=1)
add_body(doc, "本拡張機能によるユーザーデータの使用は、Chrome ウェブストア ユーザーデータ ポリシー（Limited Use要件を含みます）に準拠します。ユーザーデータは、本拡張機能が明示する単一目的および利用者向け機能の提供・改善に必要な範囲でのみ使用し、パーソナライズ広告、リターゲティング広告、信用力判断または融資目的には使用しません。")

doc.add_heading("11. ポリシーの変更", level=1)
add_body(doc, "機能またはデータ取扱いを変更する場合、本ポリシーを更新します。ユーザーデータの取扱いに重要な変更がある場合は、適用前に拡張機能内またはChromeウェブストア掲載ページで分かりやすく通知し、必要に応じて改めて同意を取得します。")

doc.add_heading("12. お問い合わせ", level=1)
add_body(doc, "本ポリシーまたは本拡張機能に関するお問い合わせは、本拡張機能のChromeウェブストア掲載ページに記載されたサポート窓口からご連絡ください。")

doc.core_properties.title = "めっちゃX プライバシーポリシー"
doc.core_properties.subject = "Chrome拡張機能のユーザーデータ取扱方針"
doc.core_properties.author = "yuyudaodevlab"
doc.core_properties.keywords = "Chrome extension, privacy policy, X Post History"
doc.save(OUT)
print(OUT)
